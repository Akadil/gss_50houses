import docx
import os
from docx.shared import Pt

old_all_files = os.listdir()
all_files = []

for item in old_all_files:

    if item[:2] == '~$':
        continue

    if item[-5:] == '.docx':
        all_files.append(item)

diff_types = [' технадзор', 'кт приемки', 'ение автор', 'Декларация']
for item in all_files:
    #print(item[-15:-5])
    doc = docx.Document(item)
    paragraphs = doc.paragraphs
    if item[-15:-5] == ' технадзор':
        print("11")

    elif item[-15:-5] == 'кт приемки':
        print("9, 12")
        run = paragraphs[9].add_run()
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        run = paragraphs[12].add_run()
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
    elif item[-15:-5] == 'ение автор':
        print("I am here")
        paragraphs[19].add_run().font.name = 'Times New Roman'
        paragraphs[19].add_run().font.size = Pt(12)

    else:
        print("13")
        run = paragraphs[13].add_run()
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    doc.save(item)
